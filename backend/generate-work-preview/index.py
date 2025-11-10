import json
import os
import zipfile
import tempfile
from typing import Dict, Any, Optional, List
import requests
import boto3
from docx import Document
from PIL import Image, ImageDraw, ImageFont
import io

def handler(event: Dict[str, Any], context: Any) -> Dict[str, Any]:
    '''
    Business: Генерация превью (скриншотов страниц "Содержание" и "Введение") из Word файлов в ZIP архиве работы
    Args: event - dict с httpMethod, queryStringParameters (work_id)
          context - объект с request_id
    Returns: HTTP response с URLs сгенерированных скриншотов
    '''
    
    if event.get('httpMethod') == 'OPTIONS':
        return {
            'statusCode': 200,
            'headers': {
                'Access-Control-Allow-Origin': '*',
                'Access-Control-Allow-Methods': 'POST, OPTIONS',
                'Access-Control-Allow-Headers': 'Content-Type, X-Auth-Token',
                'Access-Control-Max-Age': '86400'
            },
            'body': '',
            'isBase64Encoded': False
        }
    
    work_id = event.get('queryStringParameters', {}).get('work_id')
    
    if not work_id:
        return {
            'statusCode': 400,
            'headers': {
                'Content-Type': 'application/json',
                'Access-Control-Allow-Origin': '*'
            },
            'body': json.dumps({'error': 'work_id required'}),
            'isBase64Encoded': False
        }
    
    try:
        import psycopg2
        database_url = os.environ.get('DATABASE_URL')
        conn = psycopg2.connect(database_url)
        cur = conn.cursor()
        
        # Получаем информацию о работе
        cur.execute(
            "SELECT title, file_url FROM t_p63326274_course_download_plat.works WHERE id = %s",
            (work_id,)
        )
        
        result = cur.fetchone()
        
        if not result:
            cur.close()
            conn.close()
            return {
                'statusCode': 404,
                'headers': {
                    'Content-Type': 'application/json',
                    'Access-Control-Allow-Origin': '*'
                },
                'body': json.dumps({'error': 'Work not found'}),
                'isBase64Encoded': False
            }
        
        title, file_url = result
        
        if not file_url:
            cur.close()
            conn.close()
            return {
                'statusCode': 400,
                'headers': {
                    'Content-Type': 'application/json',
                    'Access-Control-Allow-Origin': '*'
                },
                'body': json.dumps({'error': 'No file URL for this work'}),
                'isBase64Encoded': False
            }
        
        # Скачиваем ZIP архив
        print(f"Downloading ZIP from {file_url}")
        zip_response = requests.get(file_url, timeout=60)
        zip_response.raise_for_status()
        
        preview_urls = []
        
        # Обрабатываем ZIP архив
        with tempfile.TemporaryDirectory() as temp_dir:
            zip_path = os.path.join(temp_dir, 'work.zip')
            with open(zip_path, 'wb') as f:
                f.write(zip_response.content)
            
            # Распаковываем ZIP
            with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                zip_ref.extractall(temp_dir)
            
            # Ищем Word файлы
            docx_files = []
            for root, dirs, files in os.walk(temp_dir):
                for file in files:
                    if file.endswith('.docx') and not file.startswith('~$'):
                        docx_files.append(os.path.join(root, file))
            
            if not docx_files:
                cur.close()
                conn.close()
                return {
                    'statusCode': 404,
                    'headers': {
                        'Content-Type': 'application/json',
                        'Access-Control-Allow-Origin': '*'
                    },
                    'body': json.dumps({'error': 'No Word files found in archive'}),
                    'isBase64Encoded': False
                }
            
            # Берем первый Word файл (обычно это ПЗ)
            docx_path = docx_files[0]
            print(f"Processing Word file: {docx_path}")
            
            # Создаем превью изображение из DOCX
            preview_url = create_formatted_preview(docx_path, work_id, temp_dir)
            if preview_url:
                preview_urls.append(preview_url)
        
        # Обновляем БД с первым скриншотом
        if preview_urls:
            cur.execute(
                "UPDATE t_p63326274_course_download_plat.works SET preview_image_url = %s WHERE id = %s",
                (preview_urls[0], work_id)
            )
            conn.commit()
        
        cur.close()
        conn.close()
        
        return {
            'statusCode': 200,
            'headers': {
                'Content-Type': 'application/json',
                'Access-Control-Allow-Origin': '*'
            },
            'body': json.dumps({
                'preview_urls': preview_urls,
                'pages_found': len(preview_urls),
                'work_id': work_id
            }),
            'isBase64Encoded': False
        }
        
    except Exception as e:
        print(f"Error generating preview: {repr(e)}")
        import traceback
        traceback.print_exc()
        return {
            'statusCode': 500,
            'headers': {
                'Content-Type': 'application/json',
                'Access-Control-Allow-Origin': '*'
            },
            'body': json.dumps({'error': 'Failed to generate preview', 'details': str(e)}),
            'isBase64Encoded': False
        }


def create_formatted_preview(docx_path: str, work_id: str, temp_dir: str) -> Optional[str]:
    """Создает красиво отформатированное превью из DOCX"""
    try:
        doc = Document(docx_path)
        
        # Ищем содержание
        content_lines = []
        found_content = False
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            text_lower = text.lower()
            
            # Нашли начало содержания
            if not found_content and any(kw in text_lower for kw in ['содержание', 'оглавление']):
                found_content = True
                content_lines.append(('title', text))
                continue
            
            # Собираем строки содержания
            if found_content:
                if text and len(text) > 2:
                    # Проверяем, не начался ли новый раздел (введение и т.д.)
                    if any(kw in text_lower for kw in ['введение', 'глава', 'раздел', 'список']):
                        if 'введение' in text_lower:
                            content_lines.append(('title', text))
                        break
                    
                    # Это строка содержания
                    content_lines.append(('item', text))
                
                # Ограничиваем количество строк
                if len(content_lines) > 30:
                    break
        
        if not content_lines:
            # Если не нашли содержание, берем первые 20 строк
            for para in doc.paragraphs[:20]:
                if para.text.strip():
                    content_lines.append(('item', para.text.strip()))
        
        # Создаем изображение
        img_width = 850
        img_height = 1200
        img = Image.new('RGB', (img_width, img_height), color='white')
        draw = ImageDraw.Draw(img)
        
        # Загружаем шрифты
        try:
            font_regular = ImageFont.truetype('/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf', 13)
            font_title = ImageFont.truetype('/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf', 14)
        except:
            font_regular = ImageFont.load_default()
            font_title = ImageFont.load_default()
        
        y = 50
        x_margin = 60
        
        for line_type, text in content_lines:
            if y > img_height - 100:
                break
            
            if line_type == 'title':
                # Заголовок (Содержание, Введение)
                draw.text((img_width // 2 - 60, y), text, fill='black', font=font_title)
                y += 35
            else:
                # Строка содержания
                # Разбиваем на номер/название и номер страницы
                parts = text.rsplit('.', 1)
                
                if len(parts) == 2 and parts[1].strip().isdigit():
                    # Есть номер страницы
                    title_part = parts[0]
                    page_num = parts[1].strip()
                    
                    # Рисуем название
                    draw.text((x_margin, y), title_part, fill='black', font=font_regular)
                    
                    # Рисуем точки
                    title_width = draw.textlength(title_part, font=font_regular)
                    dots_start = x_margin + title_width + 10
                    dots_end = img_width - x_margin - 40
                    dot_spacing = 5
                    
                    for x_dot in range(int(dots_start), int(dots_end), dot_spacing):
                        draw.text((x_dot, y + 8), '.', fill='black', font=font_regular)
                    
                    # Рисуем номер страницы
                    draw.text((dots_end + 5, y), page_num, fill='black', font=font_regular)
                else:
                    # Нет номера страницы, просто текст
                    draw.text((x_margin, y), text[:90], fill='black', font=font_regular)
                
                y += 22
        
        # Водяной знак
        watermark = "ПРЕВЬЮ"
        watermark_width = draw.textlength(watermark, font=font_title)
        draw.text((img_width - watermark_width - 50, img_height - 50), watermark, fill=(220, 220, 220), font=font_title)
        
        img_path = os.path.join(temp_dir, 'preview_content.png')
        img.save(img_path, 'PNG')
        
        # Загружаем в S3
        return upload_to_s3(img_path, work_id, 0)
        
    except Exception as e:
        print(f"Error creating formatted preview: {e}")
        import traceback
        traceback.print_exc()
        return None


def upload_to_s3(image_path: str, work_id: str, page_num: int) -> Optional[str]:
    """Загружает изображение в S3 и возвращает публичный URL"""
    try:
        s3_client = boto3.client(
            's3',
            endpoint_url='https://storage.yandexcloud.net',
            aws_access_key_id=os.environ.get('YANDEX_S3_KEY_ID'),
            aws_secret_access_key=os.environ.get('YANDEX_S3_SECRET_KEY'),
            region_name='ru-central1'
        )
        
        bucket_name = 'kyra'
        object_name = f'previews/work_{work_id}_page_{page_num}.png'
        
        s3_client.upload_file(
            image_path,
            bucket_name,
            object_name,
            ExtraArgs={'ACL': 'public-read', 'ContentType': 'image/png'}
        )
        
        url = f'https://storage.yandexcloud.net/{bucket_name}/{object_name}'
        print(f"Uploaded to S3: {url}")
        return url
    except Exception as e:
        print(f"Error uploading to S3: {e}")
        import traceback
        traceback.print_exc()
        return None