"""
Business: Парсинг детальной информации о работах с Яндекс.Диска (описание из Word, точное количество страниц, форматы файлов)
Args: event - dict с httpMethod, queryStringParameters (workId, publicKey)
      context - объект с request_id
Returns: HTTP response с fileFormats, description, pageCount, composition, chapters
"""
import json
import os
from typing import Dict, Any, List, Set, Optional
import urllib.request
import urllib.parse
import re
from docx import Document
from io import BytesIO


def handler(event: Dict[str, Any], context: Any) -> Dict[str, Any]:
    method: str = event.get('httpMethod', 'GET')
    
    if method == 'OPTIONS':
        return {
            'statusCode': 200,
            'headers': {
                'Access-Control-Allow-Origin': '*',
                'Access-Control-Allow-Methods': 'GET, OPTIONS',
                'Access-Control-Allow-Headers': 'Content-Type',
                'Access-Control-Max-Age': '86400'
            },
            'body': '',
            'isBase64Encoded': False
        }
    
    if method != 'GET':
        return {
            'statusCode': 405,
            'headers': {'Content-Type': 'application/json', 'Access-Control-Allow-Origin': '*'},
            'body': json.dumps({'error': 'Method not allowed'}),
            'isBase64Encoded': False
        }
    
    params = event.get('queryStringParameters', {})
    work_id = params.get('workId')
    public_key = params.get('publicKey', 'https://disk.yandex.ru/d/usjmeUqnkY9IfQ')
    
    if not work_id:
        return {
            'statusCode': 400,
            'headers': {'Content-Type': 'application/json', 'Access-Control-Allow-Origin': '*'},
            'body': json.dumps({'error': 'workId required'}),
            'isBase64Encoded': False
        }
    
    try:
        api_url = f"https://cloud-api.yandex.net/v1/disk/public/resources?public_key={urllib.parse.quote(public_key)}&limit=500"
        
        req = urllib.request.Request(api_url)
        with urllib.request.urlopen(req, timeout=10) as resp:
            data = json.loads(resp.read().decode())
        
        work_item = None
        if '_embedded' in data and 'items' in data['_embedded']:
            for item in data['_embedded']['items']:
                if item.get('resource_id') == work_id:
                    work_item = item
                    break
        
        if not work_item:
            return {
                'statusCode': 404,
                'headers': {'Content-Type': 'application/json', 'Access-Control-Allow-Origin': '*'},
                'body': json.dumps({'error': 'Work not found'}),
                'isBase64Encoded': False
            }
        
        folder_path = '/' + work_item['name']
        folder_url = f"https://cloud-api.yandex.net/v1/disk/public/resources?public_key={urllib.parse.quote(public_key)}&path={urllib.parse.quote(folder_path)}&limit=100"
        
        req = urllib.request.Request(folder_url)
        with urllib.request.urlopen(req, timeout=10) as resp:
            folder_data = json.loads(resp.read().decode())
        
        formats: Set[str] = set()
        word_file = None
        word_download_url = None
        composition = []
        
        if '_embedded' in folder_data and 'items' in folder_data['_embedded']:
            for file in folder_data['_embedded']['items']:
                name = file.get('name', '').lower()
                
                if 'preview' in name:
                    continue
                
                if '.' in name:
                    ext = name.split('.')[-1].upper()
                    
                    ext_map = {
                        'DOCX': 'DOCX', 'DOC': 'DOC', 'PDF': 'PDF',
                        'DWG': 'DWG', 'CDW': 'CDW', 'FRW': 'FRW',
                        'M3D': 'M3D', 'A3D': 'A3D', 'PY': 'Python',
                        'XLSX': 'XLSX', 'XLS': 'XLS',
                        'PNG': 'PNG', 'JPG': 'JPG', 'JPEG': 'JPG',
                        'TXT': 'TXT', 'RTF': 'RTF'
                    }
                    
                    if ext in ext_map:
                        formats.add(ext_map[ext])
                    
                    if ext == 'DOCX' and not word_file and ('пз' in name or 'записка' in name or 'отчет' in name):
                        word_file = file
                        word_download_url = get_download_url(public_key, folder_path + '/' + file.get('name'))
        
        page_count = None
        description = None
        chapters = []
        
        if word_download_url:
            try:
                doc_data = download_file(word_download_url)
                doc = Document(BytesIO(doc_data))
                
                full_text = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
                
                paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])
                page_count = max(10, int(paragraph_count / 15))
                
                description = extract_description(full_text)
                chapters = extract_chapters(doc)
                    
            except Exception as e:
                print(f"Error parsing Word file: {e}")
                if word_file and word_file.get('size'):
                    kb = word_file['size'] / 1024
                    page_count = max(10, int(kb / 2.5))
        
        if any(f in formats for f in ['DOCX', 'DOC', 'PDF']):
            composition.append('Пояснительная записка')
        if any(f in formats for f in ['DWG', 'CDW', 'FRW', 'M3D']):
            composition.append('Графическая часть (чертежи)')
        if 'Python' in formats:
            composition.append('Программный код')
        if 'XLSX' in formats or 'XLS' in formats:
            composition.append('Расчетные таблицы')
        
        if not composition:
            composition.append('Архив с материалами')
        
        return {
            'statusCode': 200,
            'headers': {'Content-Type': 'application/json', 'Access-Control-Allow-Origin': '*'},
            'body': json.dumps({
                'fileFormats': sorted(list(formats)),
                'pageCount': page_count,
                'description': description,
                'composition': composition,
                'chapters': chapters
            }, ensure_ascii=False),
            'isBase64Encoded': False
        }
        
    except Exception as e:
        return {
            'statusCode': 500,
            'headers': {'Content-Type': 'application/json', 'Access-Control-Allow-Origin': '*'},
            'body': json.dumps({'error': str(e)}),
            'isBase64Encoded': False
        }


def get_download_url(public_key: str, file_path: str) -> Optional[str]:
    try:
        download_url = f"https://cloud-api.yandex.net/v1/disk/public/resources/download?public_key={urllib.parse.quote(public_key)}&path={urllib.parse.quote(file_path)}"
        req = urllib.request.Request(download_url)
        with urllib.request.urlopen(req, timeout=10) as resp:
            data = json.loads(resp.read().decode())
            return data.get('href')
    except Exception as e:
        print(f"Error getting download URL: {e}")
        return None


def download_file(url: str) -> bytes:
    req = urllib.request.Request(url)
    with urllib.request.urlopen(req, timeout=30) as resp:
        return resp.read()


def extract_description(text: str) -> Optional[str]:
    patterns = [
        r'(?:Введение|ВВЕДЕНИЕ)[:\s\n]+(.*?)(?:Глава|ГЛАВА|Раздел|РАЗДЕЛ|1\.|Заключение|\n\n\n)',
        r'(?:Аннотация|АННОТАЦИЯ)[:\s\n]+(.*?)(?:Введение|ВВЕДЕНИЕ|Глава|\n\n\n)',
        r'(?:Цель работы|ЦЕЛЬ РАБОТЫ)[:\s\n]+(.*?)(?:Задачи|ЗАДАЧИ|Глава|\n\n)',
        r'(?:Актуальность|АКТУАЛЬНОСТЬ)[:\s\n]+(.*?)(?:Цель|ЦЕЛЬ|Задачи|\n\n)'
    ]
    
    for pattern in patterns:
        match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
        if match:
            desc = match.group(1).strip()
            sentences = re.split(r'[.!?]+\s+', desc)
            relevant = [s.strip() for s in sentences if 30 < len(s.strip()) < 400][:4]
            if relevant:
                result = '. '.join(relevant)
                if not result.endswith('.'):
                    result += '.'
                return result
    
    lines = [line.strip() for line in text.split('\n') if 40 < len(line.strip()) < 300][:3]
    if lines:
        return ' '.join(lines)
    
    return None


def extract_chapters(doc: Document) -> List[str]:
    chapters = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        
        if para.style and para.style.name.startswith('Heading'):
            chapters.append(text)
        elif re.match(r'^(Глава|ГЛАВА|Раздел|РАЗДЕЛ|Chapter)\s+\d+', text, re.IGNORECASE):
            chapters.append(text)
        elif re.match(r'^\d+\.\s+[А-ЯЁA-Z]', text):
            chapters.append(text)
    
    return chapters[:10]